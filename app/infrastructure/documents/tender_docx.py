from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZIP_DEFLATED, BadZipFile, ZipFile

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

from app.business.agents.tender.contracts import (
    GeneratedTenderArtifact,
    TenderAnalysis,
    TenderDocument,
    TenderDocumentBlock,
    TenderDocumentResourceStats,
    TenderOutputPlan,
)
from app.business.agents.tender.errors import (
    TenderDocumentParseError,
    TenderInputError,
    TenderRenderError,
)

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class TenderDocxReader:
    """在受控临时目录中读取 DOCX，并生成可追溯的文档块。"""

    def __init__(
        self,
        *,
        max_bytes: int = 50 * 1024 * 1024,
        hard_max_bytes: int = 70 * 1024 * 1024,
        max_uncompressed_bytes: int = 250 * 1024 * 1024,
        max_zip_entries: int = 10_000,
        max_compression_ratio: float = 200.0,
    ) -> None:
        self.max_bytes = max_bytes
        self.hard_max_bytes = hard_max_bytes
        self.max_uncompressed_bytes = max_uncompressed_bytes
        self.max_zip_entries = max_zip_entries
        self.max_compression_ratio = max_compression_ratio

    def read(self, *, file_name: str, content: bytes) -> TenderDocument:
        self._validate_input(file_name, content)
        try:
            resource_stats = self._inspect_archive(content)
            with TemporaryDirectory(prefix="tender-docx-") as temp_dir:
                source_path = Path(temp_dir) / Path(file_name).name
                source_path.write_bytes(content)
                document = Document(str(source_path))
                blocks = tuple(_iter_document_blocks(document))
        except TenderInputError:
            raise
        except Exception as exc:
            raise TenderDocumentParseError(f"招标 DOCX 解析失败：{exc}") from exc

        if not blocks:
            raise TenderDocumentParseError("招标 DOCX 不包含可读取的段落或表格。")

        source_text = "\n".join(
            (
                f"[evidence_id={block.block_id}] "
                f"[paragraph_no={block.paragraph_number}] "
                f"[block_kind={block.kind}] {block.text}"
            )
            for block in blocks
            if block.text
        )
        return TenderDocument(
            file_name=Path(file_name).name,
            content=content,
            blocks=blocks,
            source_text=source_text,
            resource_stats=resource_stats,
        )

    def _validate_input(self, file_name: str, content: bytes) -> None:
        normalized_name = Path(file_name).name
        if not normalized_name or Path(normalized_name).suffix.lower() != ".docx":
            raise TenderInputError("Tender 只接受 DOCX 招标文件。")
        if not content:
            raise TenderInputError("招标文件不能为空。")
        if len(content) > self.hard_max_bytes:
            raise TenderInputError("招标文件超过服务端硬性大小限制。")
        if len(content) > self.max_bytes:
            raise TenderInputError("招标文件超过允许的大小限制。")

    def _inspect_archive(self, content: bytes):
        try:
            with ZipFile(BytesIO(content)) as archive:
                infos = archive.infolist()
                if len(infos) > self.max_zip_entries:
                    raise TenderInputError("招标 DOCX 解压文件条目超过限制。")
                total_uncompressed = sum(info.file_size for info in infos)
                if total_uncompressed > self.max_uncompressed_bytes:
                    raise TenderInputError("招标 DOCX 解压后大小超过限制。")
                ratios = [
                    info.file_size / max(info.compress_size, 1)
                    for info in infos
                    if info.file_size
                ]
                max_ratio = max(ratios, default=0.0)
                if max_ratio > self.max_compression_ratio:
                    raise TenderInputError("招标 DOCX 压缩比异常。")
                return _resource_stats(
                    compressed_bytes=len(content),
                    uncompressed_bytes=total_uncompressed,
                    file_count=len(infos),
                    max_compression_ratio=max_ratio,
                )
        except TenderInputError:
            raise
        except (BadZipFile, OSError) as exc:
            raise TenderDocumentParseError("招标 DOCX 压缩包无法读取。") from exc


class TenderDocxSkeletonRenderer:
    """把分析计划中的源文档块复制为可填写的 DOCX 骨架。"""

    _WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _W = f"{{{_WORD_NS}}}"

    def render(
        self,
        *,
        document: TenderDocument,
        analysis: TenderAnalysis,
    ) -> tuple[GeneratedTenderArtifact, ...]:
        block_map = document.block_map()
        artifacts: list[GeneratedTenderArtifact] = []
        try:
            for output in analysis.outputs:
                artifacts.append(
                    GeneratedTenderArtifact(
                        file_name=_docx_name(output.name),
                        media_type=DOCX_MEDIA_TYPE,
                        content=self._render_output(
                            source_content=document.content,
                            output=output,
                            block_map=block_map,
                        ),
                    )
                )
        except Exception as exc:
            raise TenderRenderError(f"投标骨架文件生成失败：{exc}") from exc

        if not artifacts:
            raise TenderRenderError("投标分析没有可生成的输出文件。")
        return tuple(artifacts)

    def extract_range(
        self,
        *,
        document: TenderDocument,
        start_block_id: str,
        end_block_id: str,
        output_name: str,
    ) -> GeneratedTenderArtifact:
        """Copy a confirmed body range while preserving the source DOCX package."""

        display_name = Path(output_name).stem.strip() or "bid_format"
        output = TenderOutputPlan(
            name=display_name,
            slug="bid-format-section",
            document_label=display_name,
            source_start_block_id=start_block_id,
            source_end_block_id=end_block_id,
        )
        try:
            content = self._render_output(
                source_content=document.content,
                output=output,
                block_map=document.block_map(),
            )
        except TenderRenderError:
            raise
        except Exception as exc:
            raise TenderRenderError(f"格式章节提取失败：{exc}") from exc
        return GeneratedTenderArtifact(
            file_name=_docx_name(output_name),
            media_type=DOCX_MEDIA_TYPE,
            content=content,
        )

    def _render_output(
        self,
        *,
        source_content: bytes,
        output,
        block_map: dict[str, TenderDocumentBlock],
    ) -> bytes:
        """复制源 DOCX 包，只替换 document.xml 的 body 内容。

        保留整个源压缩包可以继续复用样式、图片、页眉页脚、编号和关系文件。
        输出计划只决定 body 中要保留的连续范围，不能把源文件外的内容写入骨架。
        """

        selected_blocks = _select_output_blocks(output=output, block_map=block_map)
        with ZipFile(BytesIO(source_content), "r") as source_zip:
            members = {
                info.filename: source_zip.read(info.filename)
                for info in source_zip.infolist()
            }

        document_xml = members.get("word/document.xml")
        if document_xml is None:
            raise TenderRenderError("源 DOCX 缺少 word/document.xml。")

        parser = etree.XMLParser(remove_blank_text=False, remove_comments=False)
        root = etree.fromstring(document_xml, parser)
        body = root.find(f".//{self._W}body")
        if body is None:
            raise TenderRenderError("源 DOCX 缺少文档正文。")

        section_properties = next(
            (child for child in body if child.tag == f"{self._W}sectPr"),
            None,
        )
        original_children = list(body)
        for child in original_children:
            body.remove(child)

        body.append(_text_paragraph(f"{output.name} - 投标骨架"))
        body.append(_text_paragraph("以下内容依据当前招标文件生成，请补充待填写内容。"))
        body.append(_page_break())

        copied_titles: set[str] = set()
        if selected_blocks:
            start_order = min(block.order for block in selected_blocks)
            end_order = max(block.order for block in selected_blocks)
            source_children = original_children[start_order - 1 : end_order]
        else:
            source_children = []

        for source_child in source_children:
            if source_child.tag == f"{self._W}sectPr":
                continue
            body.append(deepcopy(source_child))
        copied_titles.update(
            block.text.strip() for block in selected_blocks if block.text
        )

        for section_title in output.section_titles:
            normalized_title = section_title.strip()
            if not normalized_title or normalized_title in copied_titles:
                continue
            body.append(_text_paragraph(normalized_title, bold=True))
            body.append(_text_paragraph("[待填写：招标文件已明确要求此项内容。]"))

        if section_properties is not None:
            body.append(section_properties)

        members["word/document.xml"] = etree.tostring(
            root,
            encoding="UTF-8",
            xml_declaration=True,
            standalone=True,
        )
        output_buffer = BytesIO()
        with ZipFile(output_buffer, "w", compression=ZIP_DEFLATED) as output_zip:
            for member_name, content in members.items():
                output_zip.writestr(member_name, content)
        return output_buffer.getvalue()


def _iter_document_blocks(document: DocumentObject):
    heading_stack: list[str] = []
    paragraph_number = 0
    for order, child in enumerate(document.element.body.iterchildren(), start=1):
        if isinstance(child, CT_P):
            paragraph_number += 1
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            if not text:
                continue
            heading_level = _heading_level(paragraph)
            if heading_level is not None:
                heading_stack[:] = heading_stack[: heading_level - 1]
                heading_stack.append(text)
            yield TenderDocumentBlock(
                block_id=f"evidence-{order}",
                kind="paragraph",
                text=text,
                style_name=paragraph.style.name if paragraph.style else None,
                heading_level=heading_level,
                order=order,
                heading_path=tuple(heading_stack),
                is_template=_looks_like_template(text),
                paragraph_number=paragraph_number,
            )
        elif isinstance(child, CT_Tbl):
            table = Table(child, document)
            rows = tuple(
                tuple(cell.text.strip() for cell in row.cells)
                for row in table.rows
            )
            text = " | ".join(" | ".join(row) for row in rows if any(row)).strip()
            if not text:
                continue
            yield TenderDocumentBlock(
                block_id=f"evidence-{order}",
                kind="table",
                text=text,
                table_rows=rows,
                order=order,
                heading_path=tuple(heading_stack),
                table_header=rows[0] if rows else (),
                is_template=_looks_like_template(text),
            )


def _heading_level(paragraph: Paragraph) -> int | None:
    style_name = paragraph.style.name if paragraph.style else ""
    if not style_name.lower().startswith("heading"):
        return None
    suffix = style_name.removeprefix("Heading").strip()
    return int(suffix) if suffix.isdigit() else None


def _append_block(document: DocumentObject, block: TenderDocumentBlock) -> None:
    if block.kind == "table":
        table = document.add_table(rows=len(block.table_rows), cols=_max_columns(block.table_rows))
        for row_index, row in enumerate(block.table_rows):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = value
        return
    document.add_paragraph(block.text)


def _select_output_blocks(*, output, block_map: dict[str, TenderDocumentBlock]):
    start_id = output.source_start_block_id
    end_id = output.source_end_block_id
    if start_id is None or end_id is None:
        raise TenderRenderError(f"输出 {output.name} 缺少精确源文档边界。")
    start_block = block_map.get(start_id)
    end_block = block_map.get(end_id)
    if start_block is None or end_block is None:
        raise TenderRenderError(f"输出 {output.name} 引用了不存在的源文档边界。")
    start_order = start_block.order
    end_order = end_block.order
    if start_order > end_order:
        raise TenderRenderError(f"输出 {output.name} 的源文档边界顺序错误。")
    return tuple(
        block
        for block in block_map.values()
        if start_order <= block.order <= end_order
    )


def _text_paragraph(text: str, *, bold: bool = False):
    paragraph = etree.Element(f"{TenderDocxSkeletonRenderer._W}p")
    run = etree.SubElement(paragraph, f"{TenderDocxSkeletonRenderer._W}r")
    if bold:
        run_properties = etree.SubElement(run, f"{TenderDocxSkeletonRenderer._W}rPr")
        etree.SubElement(run_properties, f"{TenderDocxSkeletonRenderer._W}b")
    text_node = etree.SubElement(run, f"{TenderDocxSkeletonRenderer._W}t")
    text_node.text = text
    return paragraph


def _page_break():
    paragraph = etree.Element(f"{TenderDocxSkeletonRenderer._W}p")
    run = etree.SubElement(paragraph, f"{TenderDocxSkeletonRenderer._W}r")
    break_node = etree.SubElement(run, f"{TenderDocxSkeletonRenderer._W}br")
    break_node.set(f"{TenderDocxSkeletonRenderer._W}type", "page")
    return paragraph


def _max_columns(rows: tuple[tuple[str, ...], ...]) -> int:
    return max((len(row) for row in rows), default=1)


def _docx_name(output_name: str) -> str:
    normalized = output_name.strip() or "投标文件"
    normalized = normalized.translate(str.maketrans({char: "_" for char in r'\\/:*?"<>|'}))
    return normalized if normalized.lower().endswith(".docx") else f"{normalized}.docx"


def _looks_like_template(text: str) -> bool:
    return any(marker in text for marker in ("待填写", "模板", "格式", "盖章", "签字"))


def _resource_stats(
    *, compressed_bytes: int, uncompressed_bytes: int, file_count: int, max_compression_ratio: float
) -> TenderDocumentResourceStats:
    return TenderDocumentResourceStats(
        compressed_bytes=compressed_bytes,
        uncompressed_bytes=uncompressed_bytes,
        file_count=file_count,
        max_compression_ratio=max_compression_ratio,
    )
