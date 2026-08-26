from __future__ import annotations

from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document

from app.business.agents.tender.contracts import (
    TenderAnalysis,
    TenderOutputPlan,
)
from app.business.agents.tender.errors import TenderDocumentParseError, TenderInputError
from app.infrastructure.documents.tender_docx import (
    TenderDocxReader,
    TenderDocxSkeletonRenderer,
)


def _source_docx() -> bytes:
    document = Document()
    document.add_heading("投标文件格式", level=1)
    document.add_paragraph("投标函")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目名称"
    table.cell(0, 1).text = "待填写"
    table.cell(1, 0).text = "投标人"
    table.cell(1, 1).text = "待填写"
    document.add_paragraph("六份评估报告将在中标后分别提交。")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_tender_docx_reader_extracts_evidence_blocks_and_source_text() -> None:
    parsed = TenderDocxReader().read(file_name="招标文件.docx", content=_source_docx())

    assert len(parsed.blocks) == 4
    assert parsed.blocks[0].block_id.startswith("evidence-")
    assert parsed.blocks[2].kind == "table"
    assert "[evidence_id=" in parsed.source_text
    assert "投标函" in parsed.source_text
    assert parsed.blocks[0].heading_path == ("投标文件格式",)
    assert parsed.blocks[2].table_header == ("项目名称", "待填写")
    assert parsed.resource_stats is not None
    assert parsed.resource_stats.file_count > 0


@pytest.mark.parametrize(
    ("file_name", "content"),
    [("招标文件.pdf", b"content"), ("招标文件.docx", b""), ("招标文件.docx", b"not docx")],
)
def test_tender_docx_reader_rejects_invalid_input(file_name: str, content: bytes) -> None:
    with pytest.raises((TenderInputError, TenderDocumentParseError)):
        TenderDocxReader().read(file_name=file_name, content=content)


def test_tender_docx_reader_rejects_oversized_input() -> None:
    with pytest.raises(TenderInputError, match="大小限制"):
        TenderDocxReader(max_bytes=4).read(file_name="招标文件.docx", content=b"12345")


def test_tender_docx_reader_rejects_hard_size_limit_before_parsing() -> None:
    with pytest.raises(TenderInputError, match="硬性"):
        TenderDocxReader(max_bytes=5, hard_max_bytes=5).read(
            file_name="招标文件.docx", content=b"123456"
        )


def test_tender_docx_reader_rejects_abnormal_compression_ratio() -> None:
    buffer = BytesIO()
    with ZipFile(buffer, "w", compression=ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", "0" * 1_000)

    with pytest.raises(TenderInputError, match="压缩比"):
        TenderDocxReader(max_bytes=10_000, max_compression_ratio=2).read(
            file_name="招标文件.docx", content=buffer.getvalue()
        )


def test_tender_docx_renderer_creates_single_fillable_skeleton() -> None:
    source = TenderDocxReader().read(file_name="招标文件.docx", content=_source_docx())
    analysis = TenderAnalysis(
        status="completed",
        package_type="single_volume",
        summary="一份投标文件。",
        outputs=[
            TenderOutputPlan(
                name="投标文件",
                slug="bid",
                document_label="投标文件",
                section_titles=["投标函", "资格证明文件"],
                evidence_refs=[source.blocks[0].block_id, source.blocks[2].block_id],
                source_start_block_id=source.blocks[0].block_id,
                source_end_block_id=source.blocks[2].block_id,
            )
        ],
    )

    artifacts = TenderDocxSkeletonRenderer().render(document=source, analysis=analysis)
    rendered = Document(BytesIO(artifacts[0].content))
    paragraphs = [item.text for item in rendered.paragraphs]

    assert artifacts[0].file_name == "投标文件.docx"
    assert "投标文件格式" in paragraphs
    assert "资格证明文件" in paragraphs
    table_text = " | ".join(
        cell.text for table in rendered.tables for row in table.rows for cell in row.cells
    )
    assert "待填写" in table_text
    assert "六份评估报告将在中标后分别提交。" not in paragraphs


def test_tender_docx_renderer_preserves_source_package_parts_and_xml_formatting() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("保留原始格式")
    run.bold = True
    document.add_paragraph("格式区域结束")
    source_buffer = BytesIO()
    document.save(source_buffer)
    source_content = source_buffer.getvalue()

    source = TenderDocxReader().read(file_name="格式模板.docx", content=source_content)
    analysis = TenderAnalysis(
        status="completed",
        package_type="single_volume",
        summary="保留原始 DOCX 包。",
        outputs=[
            TenderOutputPlan(
                name="格式骨架",
                slug="format",
                    document_label="格式骨架",
                    evidence_refs=[source.blocks[0].block_id, source.blocks[1].block_id],
                    source_start_block_id=source.blocks[0].block_id,
                    source_end_block_id=source.blocks[1].block_id,
                )
        ],
    )

    artifact = TenderDocxSkeletonRenderer().render(
        document=source,
        analysis=analysis,
    )[0]

    with ZipFile(BytesIO(source_content)) as source_zip:
        with ZipFile(BytesIO(artifact.content)) as rendered_zip:
            source_names = {item.filename for item in source_zip.infolist()}
            rendered_names = {item.filename for item in rendered_zip.infolist()}
            assert source_names <= rendered_names
            assert rendered_zip.read("word/styles.xml") == source_zip.read("word/styles.xml")
            rendered_document_xml = rendered_zip.read("word/document.xml")

    assert "保留原始格式" in Document(BytesIO(artifact.content)).paragraphs[3].text
    assert b"<w:b" in rendered_document_xml


def test_tender_docx_renderer_creates_multiple_volumes() -> None:
    source = TenderDocxReader().read(file_name="招标文件.docx", content=_source_docx())
    analysis = TenderAnalysis(
        status="completed",
        package_type="multi_volume",
        summary="分为技术标和商务标。",
        outputs=[
            TenderOutputPlan(
                name="技术标",
                slug="technical",
                document_label="技术标",
                    section_titles=["技术方案"],
                    evidence_refs=[source.blocks[0].block_id],
                    source_start_block_id=source.blocks[0].block_id,
                    source_end_block_id=source.blocks[2].block_id,
                ),
            TenderOutputPlan(
                name="商务标",
                slug="commercial",
                document_label="商务标",
                    section_titles=["商务响应"],
                    evidence_refs=[source.blocks[3].block_id],
                    source_start_block_id=source.blocks[3].block_id,
                    source_end_block_id=source.blocks[3].block_id,
                ),
        ],
    )

    artifacts = TenderDocxSkeletonRenderer().render(document=source, analysis=analysis)

    assert [artifact.file_name for artifact in artifacts] == ["技术标.docx", "商务标.docx"]
    assert all(Document(BytesIO(artifact.content)).paragraphs for artifact in artifacts)
