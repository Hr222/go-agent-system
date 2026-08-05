from __future__ import annotations

import asyncio
import base64
from dataclasses import dataclass
from io import BytesIO

from docx import Document
from mcp.types import CallToolResult, EmbeddedResource

from app.infrastructure.documents.tender_docx import (
    TenderDocxReader,
    TenderDocxSkeletonRenderer,
)
from app.interfaces.agent.tender_mcp import (
    TENDER_MCP_EXTRACT_TOOL_NAME,
    TENDER_MCP_TOOL_NAME,
    TENDER_MCP_VERIFY_TOOL_NAME,
    create_tender_mcp_server,
)
from app.modules.agent.tender.application.service import TenderApplication
from app.modules.agent.tender.contracts import (
    GeneratedTenderArtifact,
    TenderAnalysis,
    TenderBoundaryContextBlock,
    TenderExtractFormatSectionResult,
    TenderGenerateSkeletonResult,
    TenderOutputPlan,
    TenderSourceEvidence,
    TenderVerifyExtractionBoundaryResult,
)
from app.modules.llm.contracts import StructuredLlmResult


@dataclass
class FakeTenderApplication:
    result: TenderGenerateSkeletonResult
    received: object | None = None
    extracted: object | None = None
    verified: object | None = None

    def execute(self, command: object) -> TenderGenerateSkeletonResult:
        self.received = command
        return self.result

    def extract_bid_format_section(self, command: object) -> TenderExtractFormatSectionResult:
        self.extracted = command
        artifact = self.result.artifacts[0]
        return TenderExtractFormatSectionResult(
            artifact=artifact,
            start_block_id=command.start_block_id,
            end_block_id=command.end_block_id,
            block_count=2,
            table_count=1,
        )

    def verify_extraction_boundary(self, command: object) -> TenderVerifyExtractionBoundaryResult:
        self.verified = command
        return TenderVerifyExtractionBoundaryResult(
            start_block_id=command.start_block_id,
            end_block_id=command.end_block_id,
            start_position=1,
            end_position=2,
            context=(
                TenderBoundaryContextBlock(
                    block_id=command.start_block_id,
                    kind="paragraph",
                    text="format start",
                    order=2,
                    position=1,
                    heading_path=("Bid format",),
                ),
            ),
        )


def _result() -> TenderGenerateSkeletonResult:
    return TenderGenerateSkeletonResult(
        analysis=TenderAnalysis(
            status="completed",
            package_type="single_volume",
            summary="生成一份投标文件。",
            outputs=[],
        ),
        artifacts=(
            GeneratedTenderArtifact(
                file_name="投标文件.docx",
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                content=b"docx-content",
            ),
        ),
        model="fake-model",
        prompt_version="tender-skeleton-v1",
    )


def test_tender_mcp_lists_only_v1_tools() -> None:
    server = create_tender_mcp_server(FakeTenderApplication(_result()))

    tools = asyncio.run(server.list_tools())

    assert [tool.name for tool in tools] == [
        TENDER_MCP_TOOL_NAME,
        TENDER_MCP_EXTRACT_TOOL_NAME,
        TENDER_MCP_VERIFY_TOOL_NAME,
    ]
    assert "content_base64" in tools[0].inputSchema["properties"]
    assert "start_block_id" in tools[1].inputSchema["required"]
    assert "end_block_id" in tools[2].inputSchema["required"]
    assert "tender.fill_bid_content" not in [tool.name for tool in tools]


def test_tender_mcp_calls_application_and_returns_structured_resource() -> None:
    application = FakeTenderApplication(_result())
    server = create_tender_mcp_server(application)

    response = asyncio.run(
        server.call_tool(
            TENDER_MCP_TOOL_NAME,
            {
                "file_name": "招标文件.docx",
                "content_base64": base64.b64encode(b"source").decode("ascii"),
                "user_focus": "关注投标文件分线",
            },
        )
    )

    assert isinstance(response, CallToolResult)
    assert response.isError is False
    assert response.structuredContent["artifacts"][0]["file_name"] == "投标文件.docx"
    assert any(isinstance(block, EmbeddedResource) for block in response.content)
    assert application.received is not None
    assert application.received.content == b"source"
    assert application.received.user_focus == "关注投标文件分线"


def test_tender_mcp_returns_stable_error_for_invalid_base64() -> None:
    server = create_tender_mcp_server(FakeTenderApplication(_result()))

    response = asyncio.run(
        server.call_tool(
            TENDER_MCP_TOOL_NAME,
            {"file_name": "招标文件.docx", "content_base64": "not-base64"},
        )
    )

    assert isinstance(response, CallToolResult)
    assert response.isError is True
    assert response.structuredContent == {
        "error_code": "INVALID_INPUT",
        "message": "文件内容不是有效的 Base64。",
    }


def test_tender_mcp_calls_format_extraction_and_returns_resource() -> None:
    application = FakeTenderApplication(_result())
    server = create_tender_mcp_server(application)

    response = asyncio.run(
        server.call_tool(
            TENDER_MCP_EXTRACT_TOOL_NAME,
            {
                "file_name": "source.docx",
                "content_base64": base64.b64encode(b"source").decode("ascii"),
                "start_block_id": "evidence-2",
                "end_block_id": "evidence-4",
            },
        )
    )

    assert response.isError is False
    assert response.structuredContent["block_count"] == 2
    assert any(isinstance(block, EmbeddedResource) for block in response.content)
    assert application.extracted.content == b"source"


def test_tender_mcp_returns_boundary_context_without_llm_decision() -> None:
    application = FakeTenderApplication(_result())
    server = create_tender_mcp_server(application)

    response = asyncio.run(
        server.call_tool(
            TENDER_MCP_VERIFY_TOOL_NAME,
            {
                "file_name": "source.docx",
                "content_base64": base64.b64encode(b"source").decode("ascii"),
                "start_block_id": "evidence-2",
                "end_block_id": "evidence-4",
            },
        )
    )

    assert response.isError is False
    assert response.structuredContent["context"][0]["block_id"] == "evidence-2"
    assert application.verified.content == b"source"


def test_tender_mcp_exposes_streamable_http_app() -> None:
    server = create_tender_mcp_server(FakeTenderApplication(_result()))

    application = server.streamable_http_app()

    assert application is not None


def test_tender_mcp_runs_real_application_and_returns_openable_docx() -> None:
    source_document = Document()
    source_document.add_heading("投标文件格式", level=1)
    source_document.add_paragraph("投标函")
    source_buffer = BytesIO()
    source_document.save(source_buffer)
    source_content = source_buffer.getvalue()

    source = TenderDocxReader().read(
        file_name="招标文件.docx",
        content=source_content,
    )
    analysis = TenderAnalysis(
        status="completed",
        package_type="single_volume",
        summary="生成一份骨架。",
        evidence=[
            TenderSourceEvidence(
                evidence_id=source.blocks[0].block_id,
                location="投标文件格式",
                quote="投标文件格式",
            ),
            TenderSourceEvidence(
                evidence_id=source.blocks[-1].block_id,
                location="投标函",
                quote="投标函",
            ),
        ],
        outputs=[
            TenderOutputPlan(
                name="投标文件",
                slug="bid",
                document_label="投标文件",
                evidence_refs=[block.block_id for block in source.blocks],
                source_start_block_id=source.blocks[0].block_id,
                source_end_block_id=source.blocks[-1].block_id,
            )
        ],
    )

    class FakeStructuredLlm:
        def invoke(self, request: object, output_schema: object) -> StructuredLlmResult:
            return StructuredLlmResult(
                value=analysis,
                model="fake-model",
                prompt_version="fake-prompt",
            )

    application = TenderApplication(
        llm=FakeStructuredLlm(),
        reader=TenderDocxReader(),
        renderer=TenderDocxSkeletonRenderer(),
    )
    server = create_tender_mcp_server(application)

    response = asyncio.run(
        server.call_tool(
            TENDER_MCP_TOOL_NAME,
            {
                "file_name": "招标文件.docx",
                "content_base64": base64.b64encode(source_content).decode("ascii"),
            },
        )
    )

    assert response.isError is False
    resource_block = next(
        block for block in response.content if isinstance(block, EmbeddedResource)
    )
    rendered_content = base64.b64decode(resource_block.resource.blob)
    rendered = Document(BytesIO(rendered_content))
    assert "投标函" in [paragraph.text for paragraph in rendered.paragraphs]
